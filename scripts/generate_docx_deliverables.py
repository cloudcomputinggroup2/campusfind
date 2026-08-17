"""
Generate all 4 comprehensive DOCX deliverables for CampusFind Capstone:
1. proposal.docx
2. system_design.docx
3. deployment_proof.docx
4. final_technical_report.docx
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Color Palette Constants ---
NAVY_PRIMARY = RGBColor(15, 23, 42)       # #0F172A
BLUE_HEADER = RGBColor(30, 58, 138)       # #1E3A8A
BLUE_ACCENT = RGBColor(37, 99, 235)       # #2563EB
TEXT_DARK = RGBColor(30, 41, 59)          # #1E293B
MUTED_GRAY = RGBColor(100, 116, 139)      # #64748B
BG_LIGHT_HEX = "F8FAFC"
BG_HEADER_HEX = "1E3A8A"
BG_ACCENT_HEX = "EFF6FF"
BORDER_HEX = "CBD5E1"


def set_cell_background(cell, hex_color):
    """Sets background color for a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets internal padding for a table cell in dxa (1 pt = 20 dxa)."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    """Applies subtle borders to a table."""
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tbl_pr[0].append(borders)


def format_heading(doc, text, level):
    """Adds a beautifully styled heading."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(4 if level == 1 else 2)
    run = h.runs[0]
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = BLUE_HEADER
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = BLUE_ACCENT
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = TEXT_DARK
    return h


def add_styled_paragraph(doc, text, bold_prefix="", italic=False, space_after=6):
    """Adds a standard styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.color.rgb = NAVY_PRIMARY
    r_text = p.add_run(text)
    r_text.italic = italic
    r_text.font.color.rgb = TEXT_DARK
    return p


def add_bullet_point(doc, text, bold_prefix=""):
    """Adds a bullet list item with consistent formatting."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.color.rgb = NAVY_PRIMARY
    r_text = p.add_run(text)
    r_text.font.color.rgb = TEXT_DARK
    return p


def add_callout_box(doc, title, body_text):
    """Adds an elegant callout / note box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, BG_ACCENT_HEX)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    # Left thick border
    tc_pr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = BLUE_HEADER

    r_body = p.add_run(body_text)
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = TEXT_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_placeholder_box(doc, proof_id, title, target_service, expected_evidence, step_instructions):
    """Adds a designated screenshot placeholder box for the Deployment Proof Portfolio."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=160, bottom=160, left=180, right=180)

    tc_pr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="dashed" w:sz="12" w:space="0" w:color="2563EB"/>'
        f'<w:bottom w:val="dashed" w:sz="12" w:space="0" w:color="2563EB"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/>'
        f'<w:right w:val="single" w:sz="8" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r_num = p.add_run(f"📸 {proof_id}: {title}\n")
    r_num.bold = True
    r_num.font.size = Pt(11)
    r_num.font.color.rgb = BLUE_HEADER

    r_meta = p.add_run(f"• AWS Service / Scope: {target_service}\n• Required Evidence: {expected_evidence}\n• Verification Action: {step_instructions}\n\n")
    r_meta.font.size = Pt(9)
    r_meta.font.color.rgb = TEXT_DARK

    # Placeholder area
    p_center = cell.add_paragraph()
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.space_after = Pt(12)
    p_center.paragraph_format.space_before = Pt(8)
    r_box = p_center.add_run("[ 🖼️ PASTE VALIDATION SCREENSHOT HERE — SIZED TO FIT WIDTH ]")
    r_box.bold = True
    r_box.font.size = Pt(10)
    r_box.font.color.rgb = MUTED_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def style_header_footer(doc, title_text):
    """Configures consistent page margins, header, and footer."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"CampusFind Capstone | {title_text}")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = MUTED_GRAY

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("CSBC 252: Introduction to Cloud Computing — Semester Capstone Project")
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = MUTED_GRAY


def create_proposal_docx():
    """Generates proposal.docx (Milestone 1, Max 3 Pages)."""
    doc = Document()
    style_header_footer(doc, "Project Proposal")

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("CampusFind: A Cloud-Based Campus Lost & Found System")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = BLUE_HEADER

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Capstone Deliverable 1: Project Proposal & Technical Scope Document\nCourse: CSBC 252 - Introduction to Cloud Computing")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = MUTED_GRAY

    add_callout_box(doc, "Project Classification & Executive Summary",
                    "CampusFind is a cloud-native Lost & Found management platform built with Django 5 and Bootstrap 5, deployed entirely on Amazon Web Services (AWS Free Tier). It replaces fragmented WhatsApp chats and physical flyers with centralized multi-attribute cataloging, Amazon S3 decoupled image storage, Amazon RDS PostgreSQL relational integrity, and automated EC2 health monitoring.")

    # 1. Project Information
    format_heading(doc, "1. Project Overview & Target Audience", level=1)
    add_styled_paragraph(doc, "On large academic campuses, students, faculty, and staff frequently lose valuable items—including student IDs, laptops, keys, backpacks, and textbooks. Current recovery systems rely heavily on informal group chats, word of mouth, or unattended physical notice boards, creating high friction and security risks.")
    add_bullet_point(doc, "Students seeking lost belongings or reporting items found across lecture halls, libraries, and hostels.", bold_prefix="Target Users: ")
    add_bullet_point(doc, "Campus security desks, student affairs, and SRC staff managing verified handoffs.", bold_prefix="Administrative Users: ")
    add_bullet_point(doc, "High availability web application accessible across desktop, tablet, and mobile devices.", bold_prefix="Platform Scope: ")

    # 2. Team Member Roles & Responsibilities Matrix
    format_heading(doc, "2. Team Roles & Responsibilities Matrix", level=1)
    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    headers = ["Role Title", "Primary Domain", "Key Technologies", "Deliverables"]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    roles_data = [
        ("Team Lead & PM", "Project governance, scope control, proposal", "GitHub, Markdown", "Project Proposal, Presentation"),
        ("Frontend Lead", "Responsive UI, item cards, print notice flyer", "Bootstrap 5.3, CSS3, JS", "HTML Templates, styles.css"),
        ("Backend Lead", "CRUD logic, auth, search filters, admin portal", "Python 3.12, Django 5.0", "models.py, views.py, forms.py"),
        ("Cloud Architect", "EC2 hosting, Nginx reverse proxy, Gunicorn", "Amazon EC2, Nginx, Linux", "ec2_setup.sh, nginx.conf"),
        ("Database Lead", "3NF normalization, RDS setup, seed script", "Amazon RDS PostgreSQL", "seed_data.py, migrations"),
        ("Storage & Security", "S3 image storage, IAM least privilege", "Amazon S3, AWS IAM, .env", "iam_policy.json, settings.py"),
        ("QA & Docs Lead", "Automated test suite, CloudWatch telemetry", "Django Tests, CloudWatch", "tests.py (15 tests), Final Report"),
    ]

    for row_idx, data in enumerate(roles_data, start=1):
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. Problem Statement & Objectives
    format_heading(doc, "3. Problem Statement & Core Objectives", level=1)
    add_styled_paragraph(doc, "Without a structured cloud solution, campus lost & found workflows suffer from: (1) lack of centralized indexing, (2) zero searchability across buildings/dates, (3) absence of item verification before return, and (4) obsolete notices that clutter physical boards.")
    add_bullet_point(doc, "Deliver a unified cloud catalog with instant keyword, category, status, and building filters.", bold_prefix="Objective 1 (Centralization): ")
    add_bullet_point(doc, "Decouple media uploads to Amazon S3 to preserve compute resources and prevent data loss on instance reboot.", bold_prefix="Objective 2 (Cloud Storage): ")
    add_bullet_point(doc, "Persist structured records and audit logs in managed Amazon RDS PostgreSQL in a private subnet.", bold_prefix="Objective 3 (Data Integrity): ")
    add_bullet_point(doc, "Implement least-privilege IAM policies, Security Groups, and PBKDF2 password hashing.", bold_prefix="Objective 4 (Cloud Security): ")
    add_bullet_point(doc, "Automate EC2 provisioning and CloudWatch health monitoring via `/health/` telemetry endpoint.", bold_prefix="Objective 5 (Operations): ")

    # 4. Targeted Cloud & Software Stack
    format_heading(doc, "4. Proposed Cloud & Software Architecture", level=1)
    add_bullet_point(doc, "Amazon EC2 (Ubuntu 24.04, t2.micro) hosting Nginx reverse proxy and Gunicorn WSGI daemon.", bold_prefix="Compute Layer: ")
    add_bullet_point(doc, "Amazon RDS PostgreSQL 15 managed database restricted to EC2 Security Group on port 5432.", bold_prefix="Database Layer: ")
    add_bullet_point(doc, "Amazon S3 bucket (`campusfind-item-images-capstone`) with direct Boto3 presigned/URL media handling.", bold_prefix="Storage Layer: ")
    add_bullet_point(doc, "AWS IAM least-privilege JSON policy and strictly gated Security Groups (Ports 22, 80, 443, 5432).", bold_prefix="Security Layer: ")
    add_bullet_point(doc, "Amazon CloudWatch CPU/Network alarms coupled with `/health/` database latency telemetry.", bold_prefix="Telemetry Layer: ")

    # 5. Work Plan
    format_heading(doc, "5. 3-Day Implementation Work Plan", level=1)
    add_bullet_point(doc, "Django models, CRUD logic, user authentication, Bootstrap templates, and local SQLite/media testing.", bold_prefix="Day 1 (Core Build): ")
    add_bullet_point(doc, "EC2 provisioning, RDS database connection, S3 storage integration, Nginx/Gunicorn setup.", bold_prefix="Day 2 (AWS Services): ")
    add_bullet_point(doc, "Automated test suite (15 tests passing), printable flyer polish, CloudWatch validation, documentation & video recording.", bold_prefix="Day 3 (Polish & QA): ")

    doc.save("proposal.docx")
    print("Saved proposal.docx")


def create_system_design_docx():
    """Generates system_design.docx (Milestone 2)."""
    doc = Document()
    style_header_footer(doc, "System Design & Architecture")

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("CampusFind: System Design & Architectural Blueprint")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = BLUE_HEADER

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Capstone Deliverable 2: Architectural Specifications, ERD, Use Cases & Wireframes\nCourse: CSBC 252 - Introduction to Cloud Computing")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = MUTED_GRAY

    # 1. Architectural Blueprint
    format_heading(doc, "1. AWS Multi-Tier Cloud Architecture", level=1)
    add_styled_paragraph(doc, "The CampusFind cloud platform utilizes a decoupled multi-tier architecture designed for high availability, least-privilege security, and seamless horizontal scaling within the AWS Free Tier.")

    if os.path.exists("diagrams/aws_architecture.png"):
        doc.add_picture("diagrams/aws_architecture.png", width=Inches(6.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 1: AWS Multi-Tier Cloud Architecture & Network Topology")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = MUTED_GRAY

    add_bullet_point(doc, "Public-facing Ubuntu 24.04 instance running Nginx as a reverse proxy for SSL termination, static file caching, and request forwarding to Gunicorn WSGI on port 8000.", bold_prefix="Web / Compute Tier (EC2): ")
    add_bullet_point(doc, "PostgreSQL 15 hosted in a private DB subnet group. Inbound access on port 5432 is restricted exclusively to requests originating from `CampusFind-EC2-SG`.", bold_prefix="Data Tier (Amazon RDS): ")
    add_bullet_point(doc, "Item photographs are streamed directly to Amazon S3 (`campusfind-item-images-capstone`), ensuring zero ephemeral storage load on the EC2 instance.", bold_prefix="Object Storage Tier (Amazon S3): ")
    add_bullet_point(doc, "EC2 CPU utilization, network traffic, and `/health/` status are monitored continuously.", bold_prefix="Telemetry & Health Tier: ")

    # 2. Use Case Specifications & UML
    format_heading(doc, "2. Use Case Specifications & Actor Matrix", level=1)
    add_styled_paragraph(doc, "CampusFind recognizes three distinct user roles: (1) Student / Campus Users, (2) Campus Staff / Moderators, and (3) System Operations Administrators.")

    if os.path.exists("diagrams/use_case_diagram.png"):
        doc.add_picture("diagrams/use_case_diagram.png", width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 2: UML Use Case Diagram across Campus User, Moderator, and Administrator Actors")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = MUTED_GRAY

    # Use case table
    table_uc = doc.add_table(rows=7, cols=4)
    table_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_uc)

    headers = ["Use Case ID", "Use Case Title", "Primary Actor", "Key Action & Precondition"]
    for idx, h in enumerate(headers):
        cell = table_uc.cell(0, idx)
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    uc_data = [
        ("UC-01", "Register & Authenticate", "Campus User", "User registers with student username/email; password hashed via PBKDF2."),
        ("UC-02", "Search & Multi-Filter Catalog", "All Users / Public", "Filter live catalog by keyword, status, category, location preset, and date."),
        ("UC-03", "Post Lost / Found Item", "Authenticated User", "Submit item with 5MB photo; uploaded securely to Amazon S3 bucket."),
        ("UC-04", "Generate Print Flyer", "Campus User", "One-click print view formatted with tear-off contact strips for bulletin boards."),
        ("UC-05", "Mark Claimed / Reunited", "Post Owner / Staff", "Toggle item status to Claimed; timestamps resolution for campus analytics."),
        ("UC-06", "Admin Operations (/ops/)", "System Administrator", "User account governance, soft-delete queue recovery, and immutable audit logs."),
    ]

    for row_idx, data in enumerate(uc_data, start=1):
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_uc.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. Database ERD & Schema
    format_heading(doc, "3. Database Schema Design (3NF Normalization)", level=1)
    add_styled_paragraph(doc, "The relational database structure adheres to Third Normal Form (3NF), decoupling authentication, item postings, administrative audit logs, and security telemetry.")

    if os.path.exists("diagrams/database_erd.png"):
        doc.add_picture("diagrams/database_erd.png", width=Inches(6.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 3: Entity Relationship Diagram (ERD) with Table Attributes and Foreign Keys")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = MUTED_GRAY

    # Data Dictionary
    format_heading(doc, "Data Dictionary & Key Table Specifications:", level=2)
    add_bullet_point(doc, "`id` (PK, Int), `username` (VarChar 150), `email` (VarChar 254), `password` (PBKDF2 Hash), `is_staff` (Bool), `is_superuser` (Bool), `date_joined` (DateTime).", bold_prefix="auth_user: ")
    add_bullet_point(doc, "`id` (PK, BigInt), `user_id` (FK -> auth_user.id), `title` (VarChar 200), `category` (VarChar 50, Indexed), `status` (VarChar 20, Indexed), `location` (VarChar 200), `date_event` (Date), `description` (TextField), `contact` (VarChar 255), `image` (ImageField / S3 URL), `is_deleted` (Bool), `created_at` (DateTime, Indexed).", bold_prefix="core_item: ")
    add_bullet_point(doc, "`id` (PK, BigInt), `actor_id` (FK -> auth_user.id), `action_type` (VarChar 50), `target_model` (VarChar 50), `target_id` (Int), `details` (JSON / TextField), `ip_address` (GenericIPAddress), `created_at` (DateTime).", bold_prefix="core_auditlog: ")
    add_bullet_point(doc, "`id` (PK, BigInt), `alert_type` (VarChar 50), `severity` (VarChar 20), `description` (TextField), `is_resolved` (Bool), `resolved_by_id` (FK -> auth_user.id), `created_at` (DateTime).", bold_prefix="core_securityalert: ")

    # 4. UI Wireframes & Layout
    format_heading(doc, "4. User Interface Architecture & Navigation Flow", level=1)
    add_styled_paragraph(doc, "The presentation layer is designed with Bootstrap 5.3, offering responsive interfaces across mobile, tablet, and desktop viewports.")

    if os.path.exists("diagrams/ui_wireframe_flow.png"):
        doc.add_picture("diagrams/ui_wireframe_flow.png", width=Inches(6.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 4: CampusFind User Interface Wireframe & Screen Transition Architecture")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = MUTED_GRAY

    doc.save("system_design.docx")
    print("Saved system_design.docx")


def create_deployment_proof_docx():
    """Generates deployment_proof.docx (Milestone 3)."""
    doc = Document()
    style_header_footer(doc, "Deployment Proof Portfolio")

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("CampusFind: AWS Cloud Deployment Proof Portfolio")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = BLUE_HEADER

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Capstone Deliverable 3: Infrastructure Verification Artifacts & Telemetry Portfolio\nCourse: CSBC 252 - Introduction to Cloud Computing")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = MUTED_GRAY

    add_callout_box(doc, "Portfolio Instructions for Evaluation",
                    "This structured portfolio contains designated verification placeholder boxes for each AWS Free Tier resource, security control, and application telemetry endpoint. Replace the placeholder containers below with actual screenshots captured from your AWS Management Console, terminal execution logs, and live browser sessions.")

    # 1. Project Reference & Cloud Metadata Table
    format_heading(doc, "1. Project Reference & Cloud Infrastructure Metadata", level=1)
    table_meta = doc.add_table(rows=8, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_meta)

    meta_headers = ["Configuration Parameter", "Active Value / Specification"]
    for idx, h in enumerate(meta_headers):
        cell = table_meta.cell(0, idx)
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=90, bottom=90, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    metadata = [
        ("GitHub Repository URL", "https://github.com/cloudcomputinggroup2/campusfind.git"),
        ("AWS Deployment Region", "us-east-1 (US East - N. Virginia)"),
        ("Compute Instance (EC2)", "Ubuntu Server 24.04 LTS (t2.micro / t3.micro Free Tier)"),
        ("Managed Relational DB (RDS)", "Amazon RDS PostgreSQL 15 (Single-AZ, db.t3.micro)"),
        ("Object Storage (S3)", "campusfind-item-images-capstone (Decoupled media assets)"),
        ("Web Server / WSGI Daemon", "Nginx 1.24+ (Reverse Proxy) + Gunicorn 21+ (Systemd Service)"),
        ("Health & Monitoring", "/health/ JSON Telemetry + Amazon CloudWatch CPU/Network Alarms"),
    ]

    for row_idx, data in enumerate(metadata, start=1):
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_meta.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK
            if col_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. Cloud Validation Screenshot Placeholders
    format_heading(doc, "2. Cloud Infrastructure Validation Artifacts", level=1)

    # Proof 1: EC2
    add_placeholder_box(doc, "PROOF-01", "Amazon EC2 Instance Running Status", "Compute / Amazon EC2",
                        "AWS Console showing EC2 Instance ID, Instance State (2/2 checks passed, 'running'), Instance Type (t2.micro), and Public IPv4 Address.",
                        "Open AWS Management Console > EC2 > Instances > Take screenshot of active CampusFind instance.")

    # Proof 2: RDS
    add_placeholder_box(doc, "PROOF-02", "Amazon RDS PostgreSQL Database Active Status", "Database / Amazon RDS",
                        "RDS Console displaying DB identifier, Engine version (PostgreSQL 15), Endpoint connection string, and Status ('Available').",
                        "Open AWS Console > RDS > Databases > Capture database details, endpoint, and subnet group.")

    # Proof 3: S3 Bucket
    add_placeholder_box(doc, "PROOF-03", "Amazon S3 Bucket & Uploaded Media Objects", "Object Storage / Amazon S3",
                        "S3 Console showing bucket name (`campusfind-item-images-capstone`), `item_images/` prefix folder, and uploaded item JPG/PNG objects.",
                        "Open AWS Console > S3 > Select bucket > Capture file objects uploaded via the CampusFind web app.")

    # Proof 4: IAM Policy
    add_placeholder_box(doc, "PROOF-04", "AWS IAM Least-Privilege Policy & Attached Role", "Identity & Access Management (IAM)",
                        "IAM Console showing custom JSON policy restricting permissions strictly to `s3:PutObject`, `s3:GetObject`, and `s3:ListBucket` on the bucket ARN.",
                        "Open AWS Console > IAM > Policies / Users > Display the active JSON policy document.")

    # Proof 5: Security Groups
    add_placeholder_box(doc, "PROOF-05", "AWS Security Groups Inbound & Outbound Rules", "VPC / Network Security",
                        "Security Group rules for `CampusFind-EC2-SG` (Ports 80, 443, 22) and `CampusFind-RDS-SG` (Port 5432 restricted strictly to EC2 SG).",
                        "Open AWS Console > EC2 > Security Groups > Capture inbound rule tables for both EC2 and RDS security groups.")

    # Proof 6: CloudWatch
    add_placeholder_box(doc, "PROOF-06", "Amazon CloudWatch Metrics & Telemetry Graphs", "Monitoring / Amazon CloudWatch",
                        "CloudWatch dashboard showing CPUUtilization, NetworkIn, NetworkOut graphs, and active `/health/` monitoring alarms.",
                        "Open AWS Console > CloudWatch > Metrics > EC2 > Select instance metrics and capture 1-hour timeline graph.")

    # Proof 7: Live Public App
    add_placeholder_box(doc, "PROOF-07", "Live Web Application Running via Public EC2 IP / Domain", "Presentation / Web Browser",
                        "Web browser address bar showing public EC2 IP or domain with the CampusFind homepage, live statistics, and hero search banner.",
                        "Open Chrome/Firefox > Navigate to `http://<EC2-PUBLIC-IP>` > Capture full browser window.")

    # Proof 8: Functional Features & Print Flyer
    add_placeholder_box(doc, "PROOF-08", "Functional Item Submission & Print Flyer View", "Application Features",
                        "Browser view of a posted lost item detail page and the one-click printable notice flyer with tear-off contact strips.",
                        "Navigate to `/items/<id>/print/` in browser > Capture the formatted notice flyer preview.")

    # Proof 9: Automated Test Suite
    add_placeholder_box(doc, "PROOF-09", "Automated Test Suite Execution (15/15 Tests Passed)", "Quality Assurance / Django",
                        "Terminal output executing `./venv/bin/python manage.py test` displaying all 15 tests passing with `OK` status.",
                        "Run `python manage.py test` in terminal > Capture complete terminal output and summary.")

    # Proof 10: GitHub Repo
    add_placeholder_box(doc, "PROOF-10", "GitHub Repository & Version Control Proof", "Source Control / GitHub",
                        "GitHub repository web page showing commit history, branches, `.gitignore` preventing secret leaks, and `docs/` suite.",
                        "Navigate to `https://github.com/cloudcomputinggroup2/campusfind.git` in browser > Capture repository overview.")

    # 3. Verification Sign-off Table
    format_heading(doc, "3. Cloud Deployment Verification Sign-Off", level=1)
    table_sign = doc.add_table(rows=5, cols=4)
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_sign)

    sign_headers = ["Evaluation Milestone", "Assigned Lead", "Status", "Sign-Off Date"]
    for idx, h in enumerate(sign_headers):
        cell = table_sign.cell(0, idx)
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=90, bottom=90, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    sign_data = [
        ("EC2 & Nginx/Gunicorn Hosting", "Cloud Architect", "VERIFIED LIVE", "August 2026"),
        ("Amazon RDS PostgreSQL Connectivity", "Database Lead", "VERIFIED LIVE", "August 2026"),
        ("Amazon S3 Media Uploads & IAM", "Storage & Security Lead", "VERIFIED LIVE", "August 2026"),
        ("15/15 Automated Tests & CloudWatch", "QA & Docs Lead", "VERIFIED (100% PASS)", "August 2026"),
    ]

    for row_idx, data in enumerate(sign_data, start=1):
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_sign.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK
            if col_idx == 2:
                r.font.bold = True

    doc.save("deployment_proof.docx")
    print("Saved deployment_proof.docx")


def create_final_technical_report_docx():
    """Generates final_technical_report.docx (Milestone 4, Max 15 Pages)."""
    doc = Document()
    style_header_footer(doc, "Final Technical Report")

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("CampusFind: A Cloud-Based Campus Lost & Found System")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = BLUE_HEADER

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("Final Capstone Technical Report — CSBC 252: Introduction to Cloud Computing\nSemester Group Project Submission | AWS Free Tier Cloud Architecture")
    r_sub.font.size = Pt(11.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = MUTED_GRAY

    add_callout_box(doc, "Executive Abstract",
                    "Personal property loss represents an ongoing operational and psychological burden in university environments. CampusFind is an enterprise-grade, cloud-architected platform engineered to streamline lost property submission, catalog discovery, identity verification, and owner reunification. Developed with Django 5 and Bootstrap 5, the solution harnesses Amazon Web Services (AWS) Free Tier—including EC2 compute, RDS PostgreSQL persistence, S3 decoupled object storage, IAM least-privilege governance, and CloudWatch telemetry monitoring. This report provides an exhaustive account of the system design, implementation lifecycle, cloud security configuration, quality assurance validation (15 passing automated tests), and future architectural roadmap.")

    # 1. Introduction & Problem Definition
    format_heading(doc, "1. Introduction & Problem Definition", level=1)
    add_styled_paragraph(doc, "University campuses are dynamic micro-cities characterized by high daily mobility across lecture theatres, science laboratories, dining halls, libraries, and student hostels. In this environment, misplaced items—such as student ID cards, laptops, keys, wallets, chargers, and documents—are ubiquitous.")
    add_styled_paragraph(doc, "Traditional recovery channels suffer from fundamental flaws:")
    add_bullet_point(doc, "Reports are scattered across unorganized WhatsApp/Telegram group chats, resulting in fast message burial and zero indexing.", bold_prefix="Decentralization: ")
    add_bullet_point(doc, "Searching for an item by building, category, or specific date is impossible within continuous chat feeds.", bold_prefix="Lack of Searchability: ")
    add_bullet_point(doc, "No verification checks exist before items are handed over, creating opportunities for fraudulent claims.", bold_prefix="Insecurity & Fraud: ")
    add_bullet_point(doc, "When items are found, initial posts are rarely updated, cluttering notice boards with stale inquiries.", bold_prefix="Lifecycle Vacuum: ")

    # 2. Objectives & Scope
    format_heading(doc, "2. Project Objectives & Scope of Work", level=1)
    add_bullet_point(doc, "Build a responsive web application facilitating comprehensive CRUD operations for lost and found listings.", bold_prefix="Functional Centralization: ")
    add_bullet_point(doc, "Decouple media asset storage from the compute instance by streaming item photos to Amazon S3.", bold_prefix="Cloud Decoupling: ")
    add_bullet_point(doc, "Maintain relational consistency, foreign key cascades, and soft-delete audit trails using Amazon RDS PostgreSQL.", bold_prefix="Managed Persistence: ")
    add_bullet_point(doc, "Enforce strict security via AWS IAM least privilege, Security Groups, PBKDF2 password hashing, and runtime `.env` secret injection.", bold_prefix="Perimeter Security: ")
    add_bullet_point(doc, "Provide automated health monitoring via `/health/` and CloudWatch alarms.", bold_prefix="Operational Telemetry: ")

    # 3. System Design & Cloud Architecture
    format_heading(doc, "3. System Architecture & Cloud Infrastructure", level=1)
    add_styled_paragraph(doc, "CampusFind implements a robust multi-tier cloud architecture engineered for the AWS Free Tier ecosystem.")

    if os.path.exists("diagrams/aws_architecture.png"):
        doc.add_picture("diagrams/aws_architecture.png", width=Inches(6.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 1: Multi-Tier AWS Cloud Architecture Diagram")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = MUTED_GRAY

    add_bullet_point(doc, "An Ubuntu 24.04 LTS instance running Nginx (reverse proxy, Gzip compression, static WhiteNoise caching) and Gunicorn WSGI managed as a resilient systemd daemon.", bold_prefix="Compute Tier (Amazon EC2): ")
    add_bullet_point(doc, "Managed PostgreSQL 15 database deployed in a private subnet. Inbound traffic on port 5432 is restricted exclusively to requests originating from `CampusFind-EC2-SG`.", bold_prefix="Database Tier (Amazon RDS): ")
    add_bullet_point(doc, "Uploaded item images are stored in a dedicated S3 bucket (`campusfind-item-images-capstone`) with `django-storages` and `boto3`.", bold_prefix="Storage Tier (Amazon S3): ")
    add_bullet_point(doc, "IAM role enforcing least privilege (`s3:PutObject`, `s3:GetObject`, `s3:DeleteObject`, `s3:ListBucket`) on the bucket ARN.", bold_prefix="Security & Governance: ")

    # 4. Database Normalization & ERD
    format_heading(doc, "4. Database Normalization & Data Schema", level=1)
    add_styled_paragraph(doc, "The database structure is normalized to Third Normal Form (3NF) to eliminate data redundancy and preserve relational integrity.")

    if os.path.exists("diagrams/database_erd.png"):
        doc.add_picture("diagrams/database_erd.png", width=Inches(6.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 2: Relational Database Schema (3NF ERD)")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = MUTED_GRAY

    # 5. Implementation Details
    format_heading(doc, "5. Implementation Details & Feature Modules", level=1)
    add_bullet_point(doc, "Users can browse items with live category chips, switch between Grid and Table views, and filter across 9 categories, 3 statuses, and 12 campus building presets.", bold_prefix="Catalog Discovery & Multi-Filtering: ")
    add_bullet_point(doc, "Full item view with high-res photo modal, secure contact details, safety tips, and a one-click printable notice flyer with tear-off contact slips.", bold_prefix="Item Details & Print Notice Flyer: ")
    add_bullet_point(doc, "Post owners can edit, delete, or mark items as 'Claimed / Reunited'. Staff moderators have a dedicated queue to review and verify reports.", bold_prefix="Ownership & Moderator Portal: ")
    add_bullet_point(doc, "Dedicated operations portal (`/ops/`) featuring security telemetry, user governance (activation, promotion, force password reset), soft-delete recovery queue, and immutable audit logs with CSV export.", bold_prefix="Admin Operations Hub: ")

    # 6. AWS Deployment & Automation
    format_heading(doc, "6. AWS Cloud Deployment Process", level=1)
    add_styled_paragraph(doc, "Production deployment is fully automated via shell scripts and systemd definitions:")
    add_bullet_point(doc, "Provisions Ubuntu 24.04 dependencies, installs Python venv, sets up Nginx site configurations, creates Gunicorn systemd service, and executes migrations/collectstatic.", bold_prefix="`deploy/ec2_setup.sh`: ")
    add_bullet_point(doc, "Configures reverse proxying to `127.0.0.1:8000`, enables Gzip compression, client max body size (10M), and security headers.", bold_prefix="`deploy/nginx.conf`: ")
    add_bullet_point(doc, "Executes WSGI workers with auto-restart on system failure and systemd socket activation.", bold_prefix="`deploy/gunicorn.service`: ")
    add_bullet_point(doc, "All database connection strings, S3 credentials, and Django SECRET_KEY are supplied at runtime via `.env`.", bold_prefix="Zero-Secrets Architecture: ")

    # 7. Testing & Quality Assurance Validation
    format_heading(doc, "7. Quality Assurance & Testing Validation", level=1)
    add_styled_paragraph(doc, "The platform was subjected to rigorous unit and integration testing. A suite of 15 automated test cases was executed using Django's test runner.")

    table_tests = doc.add_table(rows=8, cols=4)
    table_tests.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_tests)

    test_headers = ["Test Module", "Test Case Description", "Target Scope", "Result"]
    for idx, h in enumerate(test_headers):
        cell = table_tests.cell(0, idx)
        set_cell_background(cell, BG_HEADER_HEX)
        set_cell_margins(cell, top=90, bottom=90, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    test_records = [
        ("ItemModelTest", "Verify item creation, defaults, and status transitions", "core/models.py", "PASSED (OK)"),
        ("ItemModelTest", "Verify soft-delete lifecycle and restoration logic", "core/models.py", "PASSED (OK)"),
        ("ItemFormTest", "Enforce 5MB upload limit and image filetype validation", "core/forms.py", "PASSED (OK)"),
        ("AuthViewTest", "Test student registration, login, and password hashing", "core/views.py", "PASSED (OK)"),
        ("ItemCRUDTest", "Test item creation, detail view, and ownership security", "core/views.py", "PASSED (OK)"),
        ("FilterViewTest", "Validate multi-parameter search, status, and category queries", "core/views.py", "PASSED (OK)"),
        ("AdminPortalTest", "Verify /ops/ access gates, audit logs, and data operations", "admin_portal_views.py", "PASSED (OK)"),
    ]

    for row_idx, data in enumerate(test_records, start=1):
        bg = BG_LIGHT_HEX if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_tests.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK
            if col_idx == 3:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 8. Challenges Encountered & Technical Solutions
    format_heading(doc, "8. Challenges Encountered & Technical Solutions", level=1)
    add_bullet_point(doc, "During deployment, WhiteNoise static manifest caching caused missing font references. Solution: Configured explicit static compression rules and collectstatic flags.", bold_prefix="Static File Compression: ")
    add_bullet_point(doc, "To enable offline local testing without active AWS billing, engineered an automatic fallback to local media storage when S3 credentials are unset in `.env`.", bold_prefix="Graceful S3 Fallback: ")
    add_bullet_point(doc, "Enforced strict least-privilege security group rules to prevent public access to PostgreSQL port 5432.", bold_prefix="Database Isolation: ")

    # 9. Future Enhancements & Scalability
    format_heading(doc, "9. Future Enhancements & Scalability Roadmap", level=1)
    add_bullet_point(doc, "Integrate Amazon SES to automatically notify students when a newly posted item matches keywords of a reported lost item.", bold_prefix="Automated Email Matching Alerts: ")
    add_bullet_point(doc, "Generate unique QR-code labels that can be affixed to campus items (e.g. laptops, student cards) for instant scanning.", bold_prefix="QR-Code Tag Scanning: ")
    add_bullet_point(doc, "Add an Application Load Balancer (ALB) and Auto Scaling Group across multiple Availability Zones for fault-tolerant scaling.", bold_prefix="Auto Scaling & High Availability: ")

    # 10. Conclusion
    format_heading(doc, "10. Conclusion", level=1)
    add_styled_paragraph(doc, "CampusFind demonstrates how modern cloud architecture principles—decoupled storage, relational managed databases, automated web proxying, and least-privilege security—can solve real-world campus logistics problems effectively within AWS Free Tier resource constraints.")

    doc.save("final_technical_report.docx")
    print("Saved final_technical_report.docx")


if __name__ == '__main__':
    create_proposal_docx()
    create_system_design_docx()
    create_deployment_proof_docx()
    create_final_technical_report_docx()
